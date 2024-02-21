import os
import pandas as pd
import re
import win32com.client as client
from docx import Document
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog, Scrollbar, Text

# Definerer basis mappen for scriptet
BASE_DIR = os.path.dirname(os.path.realpath(__file__))
TEMPLATES_DIR = os.path.join(BASE_DIR, 'Templates')
RECIPIENTS_DIR = os.path.join(BASE_DIR, 'Recipients')
FILES_TO_SEND_DIR = os.path.join(BASE_DIR, 'FilesToSend')
SIGNATURES_DIR = os.path.join(BASE_DIR, 'Signatures')

def read_docx_template(file_path):
    try:
        doc = Document(file_path)
        subject = doc.paragraphs[0].text
        body_elements = []
        
        for para in doc.paragraphs[1:]:
            if para.text.strip():  
                if para.style.name.startswith('List'):
                    body_elements.append(f"- {para.text}\n")
                else:
                    body_elements.append(f"{para.text}\n\n")

        body = "".join(body_elements).rstrip()

        return subject, body
    except Exception as e:
        messagebox.showerror("Error", f"Error reading document: {e}")
        return None, None


def create_email(recipient, attachment_file, subject, body):
    try:
        outlook = client.Dispatch('Outlook.Application')
        message = outlook.CreateItem(0)
        message.To = recipient
        message.Subject = subject
        message.Body = body
        if attachment_file:
            message.Attachments.Add(Source=attachment_file)
        return message
    except Exception as e:
        messagebox.showerror("Error", f"Could not create email: {e}")
        return None

def select_docx_template():
    root = tk.Tk()
    root.withdraw()
    file_path = filedialog.askopenfilename(initialdir=TEMPLATES_DIR,
                                           title="Select a .docx template",
                                           filetypes=[("Word documents", "*.docx")])
    root.destroy()
    return file_path

def show_mail_preview(subject, body):
    """Shows a preview of the mail."""
    preview_window = tk.Toplevel()
    preview_window.title("Mail Preview")

    preview_text = tk.Text(preview_window, height=20, width=60, wrap='word')
    preview_text.pack(fill='both', expand=True, padx=5, pady=5)

    preview_content = f"Subject: {subject}\n\nBody:\n{body}"
    
    preview_text.insert('end', preview_content)
    preview_text.config(state='disabled')
    
    preview_window.lift()  
    preview_window.focus_force()  
    preview_window.mainloop()


def show_confirmation_dialog(email_details, subject, body):
    global user_confirmation
    confirmation_window = tk.Tk()
    confirmation_window.title("Confirmation")

    text_area = Text(confirmation_window, height=15, width=80)
    text_area.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    text_area.insert(tk.END, "Recipients and their attachments:\n\n")
    for email, attachment in email_details:
        text_area.insert(tk.END, f"{email} with attachment {os.path.basename(attachment)}\n")
    text_area.config(state=tk.DISABLED)

    scrollbar = Scrollbar(confirmation_window, orient="vertical", command=text_area.yview)
    scrollbar.pack(side=tk.LEFT, fill="y")
    text_area.config(yscrollcommand=scrollbar.set)

    def on_yes():
        global user_confirmation
        user_confirmation = True
        confirmation_window.destroy()

    def on_no():
        global user_confirmation
        user_confirmation = False
        confirmation_window.destroy()

    def on_preview():
        preview_window = tk.Toplevel(confirmation_window)
        preview_window.title("Email Preview")

        
        preview_text = tk.Text(preview_window, height=20, width=60, wrap="word")
        preview_text.pack(fill=tk.BOTH, expand=True)

        
        preview_content = f"Subject: {subject}\n\nBody:\n{body}"
        preview_text.insert(tk.END, preview_content)

        # Gj√∏r tekstfeltet i preview read-only
        #preview_text.config(state=tk.DISABLED)


    button_frame = tk.Frame(confirmation_window)
    button_frame.pack(fill=tk.X, padx=5, pady=5)

    tk.Button(button_frame, text="Yes", command=on_yes).pack(side=tk.TOP, fill=tk.X)
    tk.Button(button_frame, text="No", command=on_no).pack(side=tk.TOP, fill=tk.X)
    tk.Button(button_frame, text="Preview", command=on_preview).pack(side=tk.TOP, fill=tk.X)

    confirmation_window.lift()  
    confirmation_window.focus_force()  
    confirmation_window.mainloop()

def main():
    global user_confirmation
    user_confirmation = False  
    template_file_path = select_docx_template()
    
    if not template_file_path:
        print("No file was selected.")
        return

    subject, body = read_docx_template(template_file_path)
    if subject is None or body is None:
        print("Error loading the template. Please check the file and try again.")
        return

    df = pd.read_csv(os.path.join(RECIPIENTS_DIR, 'recipients.csv'))
    email_details = []

    for index, row in df.iterrows():
        kommune_email = row['KommuneEmail']
        attachment_files = [file for file in os.listdir(FILES_TO_SEND_DIR) if str(row['KommuneNummer']) in file]

        if attachment_files:
            attachment_path = os.path.join(FILES_TO_SEND_DIR, attachment_files[0])
            email_details.append((kommune_email, attachment_path))

    if email_details:
        show_confirmation_dialog(email_details, subject, body)
        
        if user_confirmation:  
            for email, attachment in email_details:
                try:
                    email_message = create_email(email, attachment, subject, body)
                    email_message.Send()
                    print(f"Email sent to {email} with attachment {os.path.basename(attachment)}")
                except Exception as e:
                    print(f"Could not send email to {email}: {e}")
            print("All emails have been sent.")  
        else:
            print("Sending cancelled.")  

if __name__ == "__main__":
    main()